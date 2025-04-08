from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, jsonify
import mysql.connector
from werkzeug.security import generate_password_hash, check_password_hash
import os
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
import pytesseract
import re
from docx import Document
from PIL import Image
from datetime import datetime, timedelta
import requests
import io
from docx2pdf import convert as docx2pdf_convert
import tempfile
import json
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.colors import black, blue
from docx.shared import Inches, RGBColor
from docx.oxml import parse_xml


app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Replace with a strong secret key

# Database Configuration
db_config = {
    'host': 'localhost',
    'user': 'root',  # Replace with your MySQL username
    'password': 'Furious@8',  # Replace with your MySQL password
    'database': 'db_syllabuscp'
}

@app.route('/')
def index():
    return render_template('register.html')  # Replace with your HTML file name

@app.route('/register', methods=['POST'])
def register():
    try:
        # Get form data
        first_name = request.form['FirstName']
        last_name = request.form['LastName']
        email = request.form['email']
        password = generate_password_hash(request.form['password'])  # Hash the password
        gender = request.form['gender']
        dob = request.form['dob']
        country = request.form['country']
        city = request.form['city']
        mobile_number = request.form['phno']

        # Connect to the database
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()

        # Insert data into the database
        insert_query_user = """
        INSERT INTO tbl_user (firstname, lastname, email, password, gender, dob, country, city, phno)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        insert_query_login = """
        INSERT INTO tbl_login (Username, Password, Login_type) VALUES (%s, %s, "Customer")
        """
        cursor.execute(insert_query_user, (first_name, last_name, email, password, gender, dob, country, city, mobile_number))
        cursor.execute(insert_query_login, (email, password))
        connection.commit()

        return render_template('register.html', message="User registered successfully!", status="success")

    except mysql.connector.Error as err:
        return render_template('register.html', message=str(err), status="error")

    finally:
        cursor.close()
        connection.close()

@app.route('/adminregister')
def admin_register_page():
    return render_template('adminreg.html')

@app.route('/adminregister', methods=['POST'])
def admin_register():
    try:
        # Get form data
        first_name = request.form['FirstName']
        last_name = request.form['LastName']
        email = request.form['email']
        password = generate_password_hash(request.form['password'])  # Hash the password
        gender = request.form['gender']
        dob = request.form['dob']
        country = request.form['country']
        city = request.form['city']
        mobile_number = request.form['phno']

        # Connect to the database
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()

        # Insert data into the database
        insert_query_user = """
        INSERT INTO tbl_admin (firstname, lastname, email, password, gender, dob, country, city, phno)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        insert_query_login = """
        INSERT INTO tbl_login (Username, Password, Login_type) VALUES (%s, %s, "Admin")
        """
        cursor.execute(insert_query_user, (first_name, last_name, email, password, gender, dob, country, city, mobile_number))
        cursor.execute(insert_query_login, (email, password))
        connection.commit()

        return render_template('adminreg.html', message="Admin registered successfully!", status="success")

    except mysql.connector.Error as err:
        return render_template('adminreg.html', message=str(err), status="error")

    finally:
        cursor.close()
        connection.close()

@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    try:
        username = request.form['Username']
        login_password = request.form['Login_pwd']
        
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor(dictionary=True)
        
        # Get user and login type - check both regular users and admin users
        query = """
            SELECT l.Username, l.Password, l.Login_type, 
                   COALESCE(u.UserID, a.AdminID) as UserID
            FROM tbl_login l 
            LEFT JOIN tbl_user u ON l.Username = u.email 
            LEFT JOIN tbl_admin a ON l.Username = a.email
            WHERE l.Username = %s
        """
        cursor.execute(query, (username,))
        user = cursor.fetchone()

        if user and check_password_hash(user['Password'], login_password):
            session['username'] = user['Username']
            session['user_id'] = user['UserID']
            
            # Redirect based on login type
            if user['Login_type'] == 'Admin':
                return redirect(url_for('admin_home'))
            else:
                flash('Login successful!', 'success')
                return redirect(url_for('home'))
        else:
            flash('Invalid username or password.', 'error')
            return redirect(url_for('login_page'))
    except mysql.connector.Error as err:
        flash(f"Database error: {err}", 'error')
        return redirect(url_for('login_page'))
    finally:
        cursor.close()
        connection.close()

@app.route('/home')
def home():
    if 'username' in session:
        return render_template('home.html', username=session['username'])
    else:
        flash('Please log in first.', 'error')
        return redirect(url_for('login_page'))
    
@app.route('/convert')
def convert_page():
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        cursor.execute("SELECT course_title FROM tbl_course")
        courses = [row[0] for row in cursor.fetchall()]
        print("Courses fetched:", courses)  # Debugging log
    except mysql.connector.Error as err:
        courses = []
        flash(f'Error fetching courses: {err}', 'danger')
    finally:
        cursor.close()
        connection.close()
    
    return render_template('convert.html', courses=courses)

@app.route('/extract', methods=['POST'])
def extract_text():
    user_id = session.get('user_id')  # Ensure user is logged in
    if not user_id:
        flash('Please log in to create a course plan', 'danger')
        return redirect('/login')
    
    course_title = request.form['course_title']
    start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d')
    end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d')
    
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        cursor.execute("SELECT total_hours, hours_per_week FROM tbl_course WHERE course_title = %s", (course_title,))
        course_details = cursor.fetchone()
        if course_details:
            total_hours = course_details[0]
            hours_per_week = course_details[1]
        else:
            flash('Error: Course details not found', 'danger')
            return redirect('/convert')
    except mysql.connector.Error as err:
        flash(f'Error fetching course details: {err}', 'danger')
        return redirect('/convert')
    finally:
        cursor.close()
        connection.close()
    
    units = request.form.getlist('unit_name[]')
    subtopics = request.form.getlist('unit_subtopics[]')
    unit_hours_list = request.form.getlist('unit_hours[]')

    unit_hours_list = [int(hours) for hours in unit_hours_list]

    if sum(unit_hours_list) > total_hours:
        flash('Error: Total unit hours exceed course hours.', 'danger')
        return redirect('/convert')
    
    num_units = len(units)
    total_days = (end_date - start_date).days
    hours_per_day = total_hours / total_days
    unit_hours = total_hours / num_units
    unit_days = unit_hours / hours_per_day
    
    doc = Document()
    doc.add_heading(f'Course Plan: {course_title}', level=1)
    doc.add_paragraph(f'Start Date: {start_date.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'End Date: {end_date.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Total Hours: {total_hours}')
    doc.add_paragraph(f'Hours per Week: {hours_per_week}')

    total_days = (end_date - start_date).days
    internal_exam_days = total_days // 4
    internal_exam_dates = [start_date + timedelta(days=internal_exam_days * i) for i in range(1, 4)]
    final_exam_date = end_date - timedelta(days=2)

    doc.add_paragraph('\nExam Schedule:')
    for i, exam_date in enumerate(internal_exam_dates, 1):
        doc.add_paragraph(f'Internal Exam {i}: {exam_date.strftime("%d/%m/%Y")} - {exam_date.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Final Exam: {final_exam_date.strftime("%d/%m/%Y")} - {final_exam_date.strftime("%d/%m/%Y")}')
    
    doc.add_paragraph("\n")
    doc.add_heading('Unit Details', level=2)
    
    # Create a table with better formatting
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Set table properties for better appearance
    table.autofit = True
    table.allow_autofit = True
    
    # Format header row
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        # Add bold formatting to header cells
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = 1  # Center alignment
    
    # Set header text
    hdr_cells[0].text = 'Unit'
    hdr_cells[1].text = 'Subtopics'
    hdr_cells[2].text = 'Start Date'
    hdr_cells[3].text = 'End Date'
    hdr_cells[4].text = 'Pedagogy'
    hdr_cells[5].text = 'Resources/References'
    
    # Set column widths for better content display
    table.columns[0].width = Inches(1.2)  # Unit column - increased from 1.0
    table.columns[1].width = Inches(2.8)  # Subtopics column - increased from 2.5
    table.columns[2].width = Inches(1.2)  # Start Date column - increased from 1.0
    table.columns[3].width = Inches(1.2)  # End Date column - increased from 1.0
    table.columns[4].width = Inches(1.8)  # Pedagogy column - increased from 1.5
    table.columns[5].width = Inches(1.8)  # Resources column - decreased from 2.0
    
    current_date = start_date
    for i in range(num_units):
        unit_hours = float(unit_hours_list[i])
        unit_days = int((unit_hours / (hours_per_week / 7)) + 0.5)
        end_unit_date = current_date + timedelta(days=int(unit_days))
        resources = fetch_resources(units[i])
        pedagogy = "Lecture, Discussion, Hands-on"
        
        # Add a new row
        row_cells = table.add_row().cells
        
        # Format the unit cell
        row_cells[0].text = units[i]
        for paragraph in row_cells[0].paragraphs:
            paragraph.alignment = 1  # Center alignment
        
        # Format the subtopics cell with bullet points
        subtopics_text = subtopics[i]
        if subtopics_text:
            # Split by newlines and add bullet points
            subtopics_list = [s.strip() for s in subtopics_text.split('\n') if s.strip()]
            row_cells[1].text = "• " + "\n• ".join(subtopics_list)
        
        # Format date cells
        row_cells[2].text = current_date.strftime("%d/%m/%Y")
        row_cells[3].text = end_unit_date.strftime("%d/%m/%Y")
        for j in [2, 3]:  # Center align date cells
            for paragraph in row_cells[j].paragraphs:
                paragraph.alignment = 1
        
        # Format pedagogy cell
        row_cells[4].text = pedagogy
        for paragraph in row_cells[4].paragraphs:
            paragraph.alignment = 1
        
        # Format resources cell with bullet points
        if resources:
            resources_list = [r.strip() for r in resources.split(" | ") if r.strip()]
            row_cells[5].text = "• " + "\n• ".join(resources_list)
        
        # Add alternating row colors for better readability
        if i % 2 == 1:  # Odd rows get a light gray background
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                # Use a simpler approach for cell shading
                shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
        
        current_date = end_unit_date + timedelta(days=1)
    
    # Save the document in-memory
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Store course plan in MySQL
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        insert_query = """
            INSERT INTO tbl_course_plan (UserID, plan_doc, created_at)
            VALUES (%s, %s, NOW())
        """
        cursor.execute(insert_query, (user_id, doc_io.read()))
        connection.commit()
    except mysql.connector.Error as err:
        flash(f'Error saving course plan: {err}', 'danger')
    finally:
        cursor.close()
        connection.close()
    
    doc_io.seek(0)

    return send_file(
        doc_io,
        as_attachment=True,
        download_name="course_plan.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def fetch_resources(query):
    search_engines = [
        f'https://scholar.google.com/scholar?q={query}',
        f'https://www.coursera.org/search?query={query}',
        f'https://www.edx.org/search?q={query}',
        f'https://www.khanacademy.org/search?page_search_query={query}',
        f'https://www.udemy.com/courses/search/?q={query}'
    ]
    results = []
    
    for url in search_engines:
        try:
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                results.append(url)
        except requests.exceptions.RequestException:
            continue
    
    return " | ".join(results) if results else "No references found"


@app.route('/add_course')
def add_course():
    message = session.pop('message', None)
    return render_template('course.html', message=message)

@app.route('/submit_course', methods=['POST'])
def submit_course():
    course_title = request.form['course_title']
    total_hours = request.form['total_hours']
    hours_per_week = request.form['hours_per_week']
    
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        insert_query = """
            INSERT INTO tbl_course (course_title, total_hours, hours_per_week)
            VALUES (%s, %s, %s)
        """
        cursor.execute(insert_query, (course_title, total_hours, hours_per_week))
        connection.commit()
        session['message'] = 'Course added successfully!'
    except mysql.connector.Error as err:
        session['message'] = f'Error: {err}'
    finally:
        cursor.close()
        connection.close()
    
    return redirect('/add_course')

@app.route('/course_plans')
def course_plans():
    user_id = session.get('user_id')

    if not user_id:
        flash('Please log in to view your course plans.', 'danger')
        return redirect('/login')

    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor(dictionary=True)

        cursor.execute("SELECT cp_id, created_at FROM tbl_course_plan WHERE UserID = %s ORDER BY created_at DESC", (user_id,))
        course_plans = cursor.fetchall()

    except mysql.connector.Error as err:
        flash(f"Database error: {err}", "danger")
        course_plans = []
    finally:
        cursor.close()
        connection.close()

    return render_template("course_plans.html", course_plans=course_plans)

@app.route('/download_plan/<int:plan_id>')
def download_plan(plan_id):
    user_id = session.get('user_id')

    if not user_id:
        flash('Please log in to download course plans.', 'danger')
        return redirect('/login')

    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()

        cursor.execute("SELECT plan_doc FROM tbl_course_plan WHERE cp_id = %s AND UserID = %s", (plan_id, user_id))
        result = cursor.fetchone()

        if result is None:
            flash("Course plan not found!", "danger")
            return redirect('/course_plans')

        course_plan_data = result[0]

    except mysql.connector.Error as err:
        flash(f"Error fetching file: {err}", "danger")
        return redirect('/course_plans')
    finally:
        cursor.close()
        connection.close()

    return send_file(
        io.BytesIO(course_plan_data),
        as_attachment=True,
        download_name="course_plan.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    
    return redirect(url_for('login_page'))

# Admin routes
@app.route('/admin')
def admin_dashboard():
    if 'user_id' not in session or not is_admin(session['user_id']):
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('login_page'))
    
    connection = mysql.connector.connect(**db_config)
    cursor = connection.cursor(dictionary=True)
    
    # Get dashboard statistics
    cursor.execute("SELECT COUNT(*) as total FROM tbl_user")
    total_users = cursor.fetchone()['total']
    
    cursor.execute("SELECT COUNT(*) as total FROM tbl_course_plan")
    total_plans = cursor.fetchone()['total']
    
    cursor.execute("""
        SELECT COUNT(*) as total 
        FROM tbl_user u 
        JOIN tbl_login l ON u.email = l.Username 
        WHERE l.Login_type = 'Customer'
    """)
    active_users = cursor.fetchone()['total']
    
    cursor.execute("""
        SELECT COUNT(*) as total 
        FROM tbl_course_plan 
        WHERE DATE(created_at) = CURDATE()
    """)
    today_conversions = cursor.fetchone()['total']
    
    # Get all users with their login type
    cursor.execute("""
        SELECT u.UserID, u.firstname as first_name, u.lastname as last_name, 
               u.email, l.Login_type
        FROM tbl_user u
        JOIN tbl_login l ON u.email = l.Username
    """)
    users = cursor.fetchall()
    
    # Get all course plans
    cursor.execute("""
        SELECT cp.cp_id as id, cp.created_at as created_date, 
               CONCAT(u.firstname, ' ', u.lastname) as user_name
        FROM tbl_course_plan cp
        JOIN tbl_user u ON cp.UserID = u.UserID
        ORDER BY cp.created_at DESC
    """)
    course_plans = cursor.fetchall()
    
    cursor.close()
    connection.close()
    
    return render_template('admin.html',
                         total_users=total_users,
                         total_plans=total_plans,
                         active_users=active_users,
                         today_conversions=today_conversions,
                         users=users,
                         course_plans=course_plans,
                         username=session.get('username', 'Admin'))

def is_admin(user_id):
    connection = mysql.connector.connect(**db_config)
    cursor = connection.cursor(dictionary=True)
    cursor.execute("""
        SELECT l.Login_type 
        FROM tbl_login l 
        LEFT JOIN tbl_user u ON l.Username = u.email 
        LEFT JOIN tbl_admin a ON l.Username = a.email
        WHERE (u.UserID = %s OR a.AdminID = %s) AND l.Login_type = 'Admin'
    """, (user_id, user_id))
    result = cursor.fetchone()
    cursor.close()
    connection.close()
    return result is not None

@app.route('/admin/user/<int:user_id>/edit', methods=['POST'])
def edit_user(user_id):
    if 'user_id' not in session or not is_admin(session['user_id']):
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.json
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        
        # Update user details in tbl_user
        cursor.execute("""
            UPDATE tbl_user 
            SET firstname = %s, lastname = %s, email = %s
            WHERE UserID = %s
        """, (data['first_name'], data['last_name'], data['email'], user_id))
        
        # Update login type in tbl_login
        cursor.execute("""
            UPDATE tbl_login 
            SET Login_type = %s
            WHERE Username = (SELECT email FROM tbl_user WHERE UserID = %s)
        """, (data['login_type'], user_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'message': 'User updated successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/user/<int:user_id>/delete', methods=['POST'])
def delete_user(user_id):
    if 'user_id' not in session or not is_admin(session['user_id']):
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        
        # Get user's email before deletion
        cursor.execute("SELECT email FROM tbl_user WHERE UserID = %s", (user_id,))
        user_email = cursor.fetchone()[0]
        
        # Delete user's course plans first
        cursor.execute("DELETE FROM tbl_course_plan WHERE UserID = %s", (user_id,))
        
        # Delete from tbl_login
        cursor.execute("DELETE FROM tbl_login WHERE Username = %s", (user_email,))
        
        # Delete from tbl_user
        cursor.execute("DELETE FROM tbl_user WHERE UserID = %s", (user_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'message': 'User deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/plan/<int:plan_id>/delete', methods=['POST'])
def delete_plan(plan_id):
    if 'user_id' not in session or not is_admin(session['user_id']):
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        
        cursor.execute("DELETE FROM tbl_course_plan WHERE cp_id = %s", (plan_id,))
        connection.commit()
        
        cursor.close()
        connection.close()
        
        return jsonify({'message': 'Course plan deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/settings', methods=['POST'])
def update_settings():
    if 'user_id' not in session or not is_admin(session['user_id']):
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.json
        # Here you would typically update your application settings
        # This could be stored in a database or configuration file
        return jsonify({'message': 'Settings updated successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/home')
def admin_home():
    if 'user_id' not in session or not is_admin(session['user_id']):
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('login_page'))
    
    connection = mysql.connector.connect(**db_config)
    cursor = connection.cursor(dictionary=True)
    
    # Get dashboard statistics
    cursor.execute("SELECT COUNT(*) as total FROM tbl_user")
    total_users = cursor.fetchone()['total']
    
    cursor.execute("SELECT COUNT(*) as total FROM tbl_course_plan")
    total_plans = cursor.fetchone()['total']
    
    # Get recent activities
    cursor.execute("""
        SELECT 
            cp.created_at as date,
            CONCAT(u.firstname, ' ', u.lastname) as user,
            'Created Course Plan' as action,
            'New course plan created' as details
        FROM tbl_course_plan cp
        JOIN tbl_user u ON cp.UserID = u.UserID
        ORDER BY cp.created_at DESC
        LIMIT 5
    """)
    recent_activities = cursor.fetchall()
    
    cursor.close()
    connection.close()
    
    return render_template('admin_home.html',
                         total_users=total_users,
                         total_plans=total_plans,
                         recent_activities=recent_activities)

@app.route('/edit_plan/<int:plan_id>')
def edit_plan(plan_id):
    user_id = session.get('user_id')
    if not user_id:
        flash('Please log in to edit course plans.', 'danger')
        return redirect('/login')

    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor(dictionary=True)

        # Get course plan details - modified query to not rely on course_title
        cursor.execute("""
            SELECT cp.* 
            FROM tbl_course_plan cp
            WHERE cp.cp_id = %s AND cp.UserID = %s
        """, (plan_id, user_id))
        plan = cursor.fetchone()

        if not plan:
            flash('Course plan not found!', 'danger')
            return redirect('/course_plans')

        # Parse the plan document to get units and course title
        doc = Document(io.BytesIO(plan['plan_doc']))
        units = []
        current_unit = None
        course_title = "Course Plan"  # Default value

        # Extract course title
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith('Course Plan:'):
                course_title = paragraph.text.split('Course Plan:')[1].strip()
                break

        # Extract start and end dates from the document
        start_date = None
        end_date = None
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith('Start Date:'):
                try:
                    date_str = paragraph.text.split('Start Date:')[1].strip()
                    start_date = datetime.strptime(date_str, '%d/%m/%Y')
                except:
                    start_date = datetime.now()
            elif paragraph.text.startswith('End Date:'):
                try:
                    date_str = paragraph.text.split('End Date:')[1].strip()
                    end_date = datetime.strptime(date_str, '%d/%m/%Y')
                except:
                    end_date = datetime.now() + timedelta(days=30)

        # If dates weren't found, set defaults
        if not start_date:
            start_date = datetime.now()
        if not end_date:
            end_date = start_date + timedelta(days=30)

        # Extract units from the table
        for table in doc.tables:
            if len(table.rows) > 1:  # Skip header row
                for row in table.rows[1:]:  # Skip header row
                    cells = row.cells
                    if len(cells) >= 6:  # Ensure we have enough cells
                        unit_title = cells[0].text.strip()
                        subtopics = cells[1].text.strip()
                        
                        # Extract hours from the unit title or use a default
                        unit_hours = 0
                        for paragraph in doc.paragraphs:
                            if unit_title in paragraph.text and 'hours' in paragraph.text.lower():
                                try:
                                    unit_hours = int(''.join(filter(str.isdigit, paragraph.text)))
                                except:
                                    unit_hours = 0
                                break
                        
                        # If no hours found, use a default
                        if unit_hours == 0:
                            unit_hours = 10  # Default value
                        
                        units.append({
                            'title': unit_title,
                            'subtopics': subtopics,
                            'hours': unit_hours
                        })

        # If no units were found in the table, try to extract from paragraphs
        if not units:
            current_unit = None
            for paragraph in doc.paragraphs:
                if paragraph.text.startswith('Unit'):
                    if current_unit:
                        units.append(current_unit)
                    current_unit = {
                        'title': paragraph.text.split(':', 1)[1].strip() if ':' in paragraph.text else paragraph.text,
                        'subtopics': '',
                        'hours': 0
                    }
                elif current_unit and paragraph.text.strip():
                    if 'hours' in paragraph.text.lower():
                        try:
                            current_unit['hours'] = int(''.join(filter(str.isdigit, paragraph.text)))
                        except:
                            current_unit['hours'] = 10  # Default value
                    else:
                        current_unit['subtopics'] += paragraph.text + '\n'

            if current_unit:
                units.append(current_unit)

        # If still no units, add a default one
        if not units:
            units.append({
                'title': 'Unit 1',
                'subtopics': 'Add subtopics here',
                'hours': 10
            })

        plan['course_title'] = course_title
        plan['start_date'] = start_date
        plan['end_date'] = end_date
        plan['units'] = units
        return render_template('edit_plan.html', plan=plan)

    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect('/course_plans')
    finally:
        cursor.close()
        connection.close()

@app.route('/update_plan/<int:plan_id>', methods=['POST'])
def update_plan(plan_id):
    user_id = session.get('user_id')
    if not user_id:
        flash('Please log in to update course plans.', 'danger')
        return redirect('/login')

    try:
        course_title = request.form['course_title']
        start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d')
        end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d')
        
        # Create new document
        doc = Document()
        doc.add_heading(f'Course Plan: {course_title}', level=1)
        doc.add_paragraph(f'Start Date: {start_date.strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'End Date: {end_date.strftime("%d/%m/%Y")}')
        
        # Add units
        units = request.form.getlist('unit_name[]')
        subtopics = request.form.getlist('unit_subtopics[]')
        unit_hours = request.form.getlist('unit_hours[]')
        
        # Calculate total hours and hours per week
        total_hours = sum(int(hours) for hours in unit_hours)
        weeks = (end_date - start_date).days / 7
        hours_per_week = total_hours / weeks if weeks > 0 else 0
        
        doc.add_paragraph(f'Total Hours: {total_hours}')
        doc.add_paragraph(f'Hours per Week: {hours_per_week:.1f}')
        
        # Add exam schedule
        total_days = (end_date - start_date).days
        internal_exam_days = total_days // 4
        internal_exam_dates = [start_date + timedelta(days=internal_exam_days * i) for i in range(1, 4)]
        final_exam_date = end_date - timedelta(days=2)
        
        doc.add_paragraph('\nExam Schedule:')
        for i, exam_date in enumerate(internal_exam_dates, 1):
            doc.add_paragraph(f'Internal Exam {i}: {exam_date.strftime("%d/%m/%Y")} - {exam_date.strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Final Exam: {final_exam_date.strftime("%d/%m/%Y")} - {final_exam_date.strftime("%d/%m/%Y")}')
        
        doc.add_paragraph("\n")
        doc.add_heading('Unit Details', level=2)
        
        # Create a table with better formatting
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set table properties for better appearance
        table.autofit = True
        table.allow_autofit = True
        
        # Format header row
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            # Add bold formatting to header cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = 1  # Center alignment
        
        # Set header text
        hdr_cells[0].text = 'Unit'
        hdr_cells[1].text = 'Subtopics'
        hdr_cells[2].text = 'Start Date'
        hdr_cells[3].text = 'End Date'
        hdr_cells[4].text = 'Pedagogy'
        hdr_cells[5].text = 'Resources/References'
        
        # Set column widths for better content display
        table.columns[0].width = Inches(1.2)  # Unit column - increased from 1.0
        table.columns[1].width = Inches(2.8)  # Subtopics column - increased from 2.5
        table.columns[2].width = Inches(1.2)  # Start Date column - increased from 1.0
        table.columns[3].width = Inches(1.2)  # End Date column - increased from 1.0
        table.columns[4].width = Inches(1.8)  # Pedagogy column - increased from 1.5
        table.columns[5].width = Inches(1.8)  # Resources column - decreased from 2.0
        
        current_date = start_date
        for i in range(len(units)):
            unit_hours_val = float(unit_hours[i])
            unit_days = int((unit_hours_val / (hours_per_week / 7)) + 0.5)
            end_unit_date = current_date + timedelta(days=int(unit_days))
            resources = fetch_resources(units[i])
            pedagogy = "Lecture, Discussion, Hands-on"
            
            # Add a new row
            row_cells = table.add_row().cells
            
            # Format the unit cell
            row_cells[0].text = units[i]
            for paragraph in row_cells[0].paragraphs:
                paragraph.alignment = 1  # Center alignment
            
            # Format the subtopics cell with bullet points
            subtopics_text = subtopics[i]
            if subtopics_text:
                # Split by newlines and add bullet points
                subtopics_list = [s.strip() for s in subtopics_text.split('\n') if s.strip()]
                row_cells[1].text = "• " + "\n• ".join(subtopics_list)
            
            # Format date cells
            row_cells[2].text = current_date.strftime("%d/%m/%Y")
            row_cells[3].text = end_unit_date.strftime("%d/%m/%Y")
            for j in [2, 3]:  # Center align date cells
                for paragraph in row_cells[j].paragraphs:
                    paragraph.alignment = 1
            
            # Format pedagogy cell
            row_cells[4].text = pedagogy
            for paragraph in row_cells[4].paragraphs:
                paragraph.alignment = 1
            
            # Format resources cell with bullet points
            if resources:
                resources_list = [r.strip() for r in resources.split(" | ") if r.strip()]
                row_cells[5].text = "• " + "\n• ".join(resources_list)
            
            # Add alternating row colors for better readability
            if i % 2 == 1:  # Odd rows get a light gray background
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                        # Use a simpler approach for cell shading
                        shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F2F2F2"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
            
            current_date = end_unit_date + timedelta(days=1)
        
        # Save document to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Update database
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()
        
        cursor.execute("""
            UPDATE tbl_course_plan 
            SET plan_doc = %s 
            WHERE cp_id = %s AND UserID = %s
        """, (doc_io.read(), plan_id, user_id))
        
        connection.commit()
        flash('Course plan updated successfully!', 'success')
        
    except Exception as e:
        flash(f'Error updating course plan: {str(e)}', 'danger')
    finally:
        cursor.close()
        connection.close()
    
    return redirect('/course_plans')

@app.route('/download_plan_pdf/<int:plan_id>')
def download_plan_pdf(plan_id):
    user_id = session.get('user_id')
    if not user_id:
        flash('Please log in to download course plans.', 'danger')
        return redirect('/login')

    try:
        connection = mysql.connector.connect(**db_config)
        cursor = connection.cursor()

        cursor.execute("SELECT plan_doc FROM tbl_course_plan WHERE cp_id = %s AND UserID = %s", (plan_id, user_id))
        result = cursor.fetchone()

        if result is None:
            flash("Course plan not found!", "danger")
            return redirect('/course_plans')

        # Create a BytesIO object from the DOCX data
        docx_data = io.BytesIO(result[0])
        
        # Parse the DOCX document
        doc = Document(docx_data)
        
        # Create a temporary PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
            pdf_path = pdf_file.name
        
        # Create PDF document with smaller margins to accommodate the table
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=36,  # Reduced from 72
            leftMargin=36,   # Reduced from 72
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=20,
            alignment=TA_LEFT
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        # Create cell styles for different column types
        unit_style = ParagraphStyle(
            'UnitStyle',
            parent=styles['Normal'],
            fontSize=9,
            alignment=TA_CENTER,
            spaceBefore=3,
            spaceAfter=3
        )
        
        subtopics_style = ParagraphStyle(
            'SubtopicsStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_LEFT,
            spaceBefore=3,
            spaceAfter=3,
            leading=10  # Line spacing
        )
        
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            spaceBefore=3,
            spaceAfter=3
        )
        
        pedagogy_style = ParagraphStyle(
            'PedagogyStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            spaceBefore=3,
            spaceAfter=3
        )
        
        resources_style = ParagraphStyle(
            'ResourcesStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_LEFT,
            spaceBefore=3,
            spaceAfter=3,
            leading=10  # Line spacing
        )
        
        # Build the PDF content
        story = []
        
        # Process paragraphs from the DOCX
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                story.append(Spacer(1, 12))
                continue
                
            if paragraph.style.name.startswith('Heading 1'):
                story.append(Paragraph(text, title_style))
            elif paragraph.style.name.startswith('Heading 2'):
                story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))
        
        # Process tables from the DOCX
        for table in doc.tables:
            story.append(Spacer(1, 20))
            
            # Create a table for the PDF
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            # Extract data from the DOCX table
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get text from the cell
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                data.append(row_data)
            
            # Calculate appropriate column widths based on content
            # Total width should be less than page width (letter width = 8.5 inches)
            # We'll use 7.5 inches to leave some margin
            total_width = 7.5 * inch
            
            # Define column widths as percentages of total width
            col_widths = [
                0.15 * total_width,  # Unit (15%)
                0.30 * total_width,  # Subtopics (30%)
                0.15 * total_width,  # Start Date (15%)
                0.15 * total_width,  # End Date (15%)
                0.10 * total_width,  # Pedagogy (10%)
                0.15 * total_width   # Resources (15%)
            ]
            
            # Process the data to create properly formatted cells
            formatted_data = []
            
            # Process header row
            header_row = []
            for i, cell_text in enumerate(data[0]):
                header_row.append(Paragraph(cell_text, styles['Heading3']))
            formatted_data.append(header_row)
            
            # Process data rows
            for row_idx, row in enumerate(data[1:], 1):
                formatted_row = []
                
                # Unit column (centered)
                formatted_row.append(Paragraph(row[0], unit_style))
                
                # Subtopics column (left-aligned with bullet points)
                subtopics_text = row[1]
                if subtopics_text:
                    # Split by newlines and add bullet points
                    subtopics_list = [s.strip() for s in subtopics_text.split('\n') if s.strip()]
                    if not any(s.startswith('•') for s in subtopics_list):
                        subtopics_text = "• " + "\n• ".join(subtopics_list)
                formatted_row.append(Paragraph(subtopics_text, subtopics_style))
                
                # Date columns (centered)
                formatted_row.append(Paragraph(row[2], date_style))
                formatted_row.append(Paragraph(row[3], date_style))
                
                # Pedagogy column (centered)
                formatted_row.append(Paragraph(row[4], pedagogy_style))
                
                # Resources column (left-aligned with bullet points)
                resources_text = row[5]
                if resources_text:
                    # Split by newlines and add bullet points
                    resources_list = [r.strip() for r in resources_text.split('\n') if r.strip()]
                    if not any(r.startswith('•') for r in resources_list):
                        resources_text = "• " + "\n• ".join(resources_list)
                formatted_row.append(Paragraph(resources_text, resources_style))
                
                formatted_data.append(formatted_row)
            
            # Create the PDF table with specific column widths
            pdf_table = Table(formatted_data, colWidths=col_widths)
            
            # Style the table
            table_style = TableStyle([
                # Header row styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                
                # Data rows styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                
                # Column-specific alignment
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Unit column
                ('ALIGN', (1, 1), (1, -1), 'LEFT'),    # Subtopics column
                ('ALIGN', (2, 1), (3, -1), 'CENTER'),  # Date columns
                ('ALIGN', (4, 1), (4, -1), 'CENTER'),  # Pedagogy column
                ('ALIGN', (5, 1), (5, -1), 'LEFT'),    # Resources column
                
                # Cell padding
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ])
            
            # Add alternating row colors
            for i in range(1, len(formatted_data), 2):
                table_style.add('BACKGROUND', (0, i), (-1, i), colors.lightgrey)
            
            pdf_table.setStyle(table_style)
            
            # Add the table to the story with a wrapper to handle overflow
            from reportlab.platypus import KeepTogether
            story.append(KeepTogether(pdf_table))
        
        # Build the PDF
        pdf_doc.build(story)
        
        # Send the PDF file
        response = send_file(
            pdf_path,
            as_attachment=True,
            download_name="course_plan.pdf",
            mimetype="application/pdf"
        )
        
        # Clean up temporary files
        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)
            except Exception as e:
                print(f"Error cleaning up files: {e}")
        
        return response

    except Exception as e:
        print(f"Error in download_plan_pdf: {str(e)}")
        # If PDF conversion fails, send the DOCX file instead
        try:
            return send_file(
                io.BytesIO(result[0]),
                as_attachment=True,
                download_name="course_plan.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e2:
            print(f"Error sending DOCX file: {str(e2)}")
            flash(f"Error: {str(e)}", "danger")
            return redirect('/course_plans')
    finally:
        cursor.close()
        connection.close()

@app.route('/test_pdf_conversion')
def test_pdf_conversion():
    try:
        # Create a simple document
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is a test document for PDF conversion.')
        
        # Create a temporary PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
            pdf_path = pdf_file.name
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        # Build the PDF content
        story = []
        
        # Process paragraphs from the DOCX
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                story.append(Spacer(1, 12))
                continue
                
            if paragraph.style.name.startswith('Heading 1'):
                story.append(Paragraph(text, title_style))
            else:
                story.append(Paragraph(text, normal_style))
        
        # Build the PDF
        pdf_doc.build(story)
        
        print(f"Test PDF creation successful")
        
        # Send the PDF file
        response = send_file(
            pdf_path,
            as_attachment=True,
            download_name="test.pdf",
            mimetype="application/pdf"
        )
        
        # Clean up temporary files
        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)
            except Exception as e:
                print(f"Error cleaning up files: {e}")
        
        return response
    except Exception as e:
        print(f"Test PDF creation failed: {str(e)}")
        return f"PDF creation test failed: {str(e)}"

if __name__ == '__main__':
    app.run(debug=True)
